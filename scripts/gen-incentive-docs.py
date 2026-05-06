"""봉이 인센티브 안내서 .docx 자동 생성 — 상담사용 + 팀장용
사용: python3 scripts/gen-incentive-docs.py
출력: docs/incentive-doc-agent.docx + docs/incentive-doc-manager.docx
정책: incentive_rules 활성 row 기준 (DB 직접 조회 — supabase service_role)
"""
import os
import sys
import urllib.request
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 정책 fetch (라이브 API에서) ───
def fetch_rules():
    try:
        url = 'https://bongi-mobile-production.up.railway.app/api/incentive/rules'
        req = urllib.request.Request(url, headers={'User-Agent': 'docs-gen'})
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read())
            return data.get('rule') or data.get('rules') or {}
    except Exception as e:
        print(f'[warn] rules fetch failed: {e} — defaults 사용', file=sys.stderr)
        return {}

R = fetch_rules()
def fmt(n):
    try: return f'{int(n):,}'
    except: return str(n or '0')

BASE_SAL = R.get('base_salary') or 2300000
RATES = R.get('grade_rates') or {'1':20000,'2':30000,'3':40000}
THR = R.get('grade_thresholds') or {'2':{'points':16,'premium':5},'3':{'points':31,'premium':10}}
BONUS_PER = R.get('bonus_per_premium') or 10000
PB_CO = R.get('payback_company_limit') or 30000
PB_MAX = R.get('payback_max') or 50000
MGR_RATE = R.get('manager_override_rate') or 0.12
MGR_OBLIG = R.get('manager_obligation_count') or 20
MGR_PARTIAL = R.get('manager_penalty_partial_min') or 10
MGR_PRMIN = R.get('manager_team_profit_rate_min') or 0.20
VER = R.get('version') or 'V5.1'
EFFECT = R.get('effective_from') or '2026-05-01'

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def add_table_row(table, cells, bold_first=False, header=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1e293b')
            cell._tc.get_or_add_tcPr().append(shd)
        elif bold_first and i == 0:
            run.bold = True
        run.font.size = Pt(10)

# ═══════════════════════════════════════════════
# 1) 상담사 안내서
# ═══════════════════════════════════════════════
def gen_agent():
    doc = Document()
    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

    title = doc.add_heading('👤 상담사 인센티브 안내서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f'봉이모바일 정책 {VER} · 시행일 {EFFECT}')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    add_heading(doc, '① 기본급', level=1)
    add_para(doc, f'매월 고정 지급, 영업 실적과 무관. 기본급은 {fmt(BASE_SAL)}원 (상담사별로 다를 수 있음 — 본인 급여명세서 확인).')

    add_heading(doc, '② 등급 체계 — 누적 가중치 P 기준', level=1)
    t = doc.add_table(rows=0, cols=4); t.style = 'Light Grid Accent 1'
    add_table_row(t, ['등급', '도달 조건', 'P당 단가', '페널티 조건'], header=True)
    add_table_row(t, ['G1', f'P 0 ~ {THR["2"]["points"]-1}', f'{fmt(RATES["1"])}원', '없음'], bold_first=True)
    add_table_row(t, ['G2', f'P {THR["2"]["points"]} 이상 + 우수 {THR["2"]["premium"]}건', f'{fmt(RATES["2"])}원', '우수 미달 → G1 단가 적용'], bold_first=True)
    add_table_row(t, ['G3', f'P {THR["3"]["points"]} 이상 + 우수 {THR["3"]["premium"]}건', f'{fmt(RATES["3"])}원', '우수 미달 → G2 단가 적용'], bold_first=True)
    add_para(doc, '⚠️ 페널티: 등급 P는 도달했지만 우수 상품 의무를 채우지 못하면 한 단계 낮은 등급의 단가가 적용됩니다.', size=10, color=RGBColor(0xb4,0x53,0x09))

    add_heading(doc, '③ 우수 상품 보너스', level=1)
    for line in [
        f'• 우수 상품 1건당 {fmt(BONUS_PER)}원 추가 보너스',
        '• 우수 = 회사 마진 임계값 이상 (admin이 상품별 지정)',
        '• 등급업 + 페널티 해제에 모두 활용',
    ]: add_para(doc, line, size=10)

    add_heading(doc, '④ 가중치 (Point) — 상품별 환산', level=1)
    t2 = doc.add_table(rows=0, cols=3); t2.style = 'Light Grid Accent 1'
    add_table_row(t2, ['상품 예시', '가중치', '설명'], header=True)
    add_table_row(t2, ['일반 인터넷 (500M)', '1.0 P', '기본 단위'], bold_first=True)
    add_table_row(t2, ['인터넷 1G + TV 결합', '1.5 P', '고가 결합 우대'], bold_first=True)
    add_table_row(t2, ['프리미엄 결합', '2.0 P', '최상위 상품'], bold_first=True)
    add_para(doc, '※ 정확한 가중치는 상품별 admin 설정 — 견적 시 자동 계산', size=9, color=RGBColor(0x64,0x74,0x8b))

    add_heading(doc, '⑤ 월 수령액 계산식', level=1)
    add_para(doc, '월 수령액 = 기본급 + 인센티브 + 보너스 − 본인 페이백 부담', bold=True, size=11)
    for line in [
        '   • 인센티브 = (누적 P) × (등급 단가)',
        f'   • 보너스 = (우수 건수) × {fmt(BONUS_PER)}원',
        f'   • 본인 페이백 부담 = (페이백 {fmt(PB_CO)}원 초과분)',
    ]: add_para(doc, line, size=10)

    add_heading(doc, '⑥ 계산 예시 (G2 달성)', level=1)
    add_para(doc, '조건: 영업 12건 / 누적 20P / 우수 5건 / G2 페널티 없음', size=10, color=RGBColor(0x64,0x74,0x8b))
    t3 = doc.add_table(rows=0, cols=2); t3.style = 'Light List Accent 1'
    add_table_row(t3, ['항목', '금액'], header=True)
    add_table_row(t3, ['기본급', f'{fmt(BASE_SAL)}원'])
    add_table_row(t3, ['인센티브 (20P × {fmt(RATES["2"])})'.replace('{fmt(RATES["2"])}', fmt(RATES['2'])), f'+ {fmt(20*RATES["2"])}원'])
    add_table_row(t3, [f'우수 보너스 (5건 × {fmt(BONUS_PER)})', f'+ {fmt(5*BONUS_PER)}원'])
    add_table_row(t3, [f'본인 페이백 (페이백 {fmt(PB_MAX)}원 1건 → 본인 부담)', f'− {fmt(PB_MAX-PB_CO)}원'])
    total = BASE_SAL + 20*RATES['2'] + 5*BONUS_PER - (PB_MAX-PB_CO)
    add_table_row(t3, ['총 수령액', f'{fmt(total)}원'], bold_first=True)

    add_heading(doc, '⑦ 페이백 부담 규칙', level=1)
    t4 = doc.add_table(rows=0, cols=3); t4.style = 'Light Grid Accent 1'
    add_table_row(t4, ['구간', '부담 주체', '비고'], header=True)
    add_table_row(t4, [f'0 ~ {fmt(PB_CO)}원', '회사 부담', '회사가 전액 처리'])
    add_table_row(t4, [f'{fmt(PB_CO)} ~ {fmt(PB_MAX)}원', '상담사 부담', '초과분만큼 본인 인센티브에서 차감'])
    add_table_row(t4, [f'{fmt(PB_MAX)}원 초과', '등록 불가', 'admin 승인 필요'])
    add_para(doc, f'💡 예시: 페이백 {fmt(PB_MAX)}원 영업 1건 → 회사 {fmt(PB_CO)}원 부담, 본인 {fmt(PB_MAX-PB_CO)}원 차감', size=10)

    add_heading(doc, '⑧ 등급업 전략', level=1)
    for line in [
        f'• G1 → G2: {THR["2"]["points"]}P 도달 + 우수 {THR["2"]["premium"]}건 → 단가 +50% ({fmt(RATES["1"])} → {fmt(RATES["2"])})',
        f'• G2 → G3: {THR["3"]["points"]}P 도달 + 우수 {THR["3"]["premium"]}건 → 단가 추가 +33% ({fmt(RATES["2"])} → {fmt(RATES["3"])})',
        '• 페널티 회피: P는 도달했지만 우수 부족 시 손해 큼 — 우수 등록 우선',
        '• 실시간 시뮬레이터: TM 상담 페이지에서 견적 1건 추가 시 등급 변화 확인',
    ]: add_para(doc, line, size=10)

    add_heading(doc, '⑨ 정산 시점 · 영업 상태', level=1)
    t5 = doc.add_table(rows=0, cols=3); t5.style = 'Light Grid Accent 1'
    add_table_row(t5, ['상태', '정산 포함', '설명'], header=True)
    add_table_row(t5, ['⏰ 계약대기 (pending)', '❌', 'TM 등록 직후 — 계약부서 처리 대기'])
    add_table_row(t5, ['🚧 계약진행 (in_progress)', '❌', '계약서 작성 중'])
    add_table_row(t5, ['✅ 계약완료 (completed)', '✅', '인센티브·보너스·P 모두 카운트'])
    add_table_row(t5, ['❌ 취소 (cancelled)', '❌', '고객 변심·설치 불가 등'])

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(f'\n봉이모바일 인센티브 정책 {VER} · 마지막 업데이트 {EFFECT}')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64,0x74,0x8b); r.italic = True

    out = '/Users/vinsenzo/bongi-mobile/docs/incentive-doc-agent.docx'
    doc.save(out)
    print(f'✅ {out}')

# ═══════════════════════════════════════════════
# 2) 팀장 안내서
# ═══════════════════════════════════════════════
def gen_manager():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

    title = doc.add_heading('🏢 팀장(manager) 인센티브 안내서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f'봉이모바일 정책 {VER} · 시행일 {EFFECT}')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p = doc.add_paragraph()
    r = p.add_run('📌 팀장 급여 구성\n'); r.bold = True; r.font.size = Pt(11)
    r = p.add_run('1) 본인 영업 인센티브 — V5 정책 그대로 (별첨: 상담사 안내서 참조)\n'); r.font.size = Pt(10)
    r = p.add_run('2) 팀 오버라이드 — 본 안내서의 V5.1 정책 적용 (페널티는 오버라이드에만 발동, 본인 인센티브는 영향 없음)'); r.font.size = Pt(10)

    add_heading(doc, '① 기본급', level=1)
    add_para(doc, f'매월 고정 지급. 팀장 기본급은 {fmt(BASE_SAL)}원 (개별 상담사별로 admin 설정 가능 — 본인 급여명세서 확인).')

    add_heading(doc, '② 본인 영업 인센티브 (V5)', level=1)
    add_para(doc, '등급 P × 단가 + 우수 보너스 − 페이백 부담. 자세한 계산은 상담사 인센티브 안내서 참조.')
    t = doc.add_table(rows=0, cols=3); t.style = 'Light Grid Accent 4'
    add_table_row(t, ['등급', '도달 조건', 'P당 단가'], header=True)
    add_table_row(t, ['G1', 'P 0~15', f'{fmt(RATES["1"])}원'], bold_first=True)
    add_table_row(t, ['G2', f'P {THR["2"]["points"]}+ / 우수 {THR["2"]["premium"]}건', f'{fmt(RATES["2"])}원'], bold_first=True)
    add_table_row(t, ['G3', f'P {THR["3"]["points"]}+ / 우수 {THR["3"]["premium"]}건', f'{fmt(RATES["3"])}원'], bold_first=True)

    add_heading(doc, '③ 팀 오버라이드 (Team Override) — V5.1 핵심', level=1)
    add_para(doc, f'팀 오버라이드 = 팀원 인센·보너스 합계 × {int(MGR_RATE*100)}% × 페널티 계수 (0% / 50% / 100%)', bold=True)
    for line in [
        '• 팀원 정의: 같은 센터의 role=agent + active 상담사 (본인은 제외)',
        '• 합계 기준: 각 팀원의 (인센티브 + 우수 보너스) — 본인 인센티브는 산정 제외',
        '• 지급 시점: 팀원이 영업을 \'completed\'로 등록 → 즉시 반영 (월말 정산 기다릴 필요 X)',
    ]: add_para(doc, line, size=10)

    add_heading(doc, '④ 페널티 매트릭스', level=1)
    t2 = doc.add_table(rows=0, cols=4); t2.style = 'Light Grid Accent 4'
    add_table_row(t2, ['본인 영업 (completed)', '팀 영업이익률', '페널티 계수', '오버라이드 지급'], header=True)
    add_table_row(t2, [f'≥ {MGR_OBLIG}건 (의무 충족)', f'≥ {int(MGR_PRMIN*100)}%', '100%', '정상 지급'])
    add_table_row(t2, ['≥ 의무 건수', f'< {int(MGR_PRMIN*100)}%', '0%', '미지급'])
    add_table_row(t2, [f'{MGR_PARTIAL} ~ 의무-1 건', '≥ 임계', '50%', '절반 차감'])
    add_table_row(t2, [f'{MGR_PARTIAL} ~ 의무-1 건', '< 임계', '0%', '미지급'])
    add_table_row(t2, [f'< {MGR_PARTIAL}건', '(무관)', '0%', '미지급'])

    add_heading(doc, '⑤ 팀 영업이익률 정의', level=1)
    add_para(doc, '팀 영업이익률 = (팀 매출 − 팀 페이백) ÷ 팀 매출', bold=True)
    for line in [
        '• 팀원(본인 제외) 매출·페이백 합계 기준',
        f'• {int(MGR_PRMIN*100)}% 미달 시 오버라이드 지급률 0%',
    ]: add_para(doc, line, size=10)

    add_heading(doc, '⑥ 페널티 면제 조건', level=1)
    t3 = doc.add_table(rows=0, cols=2); t3.style = 'Light Grid Accent 4'
    add_table_row(t3, ['면제 사유', '조건'], header=True)
    add_table_row(t3, ['🛡 신입 OJT', '신규 팀장 인계 기간'])
    add_table_row(t3, ['🛡 시스템 장애', '회사 측 사유로 영업 불가'])
    add_table_row(t3, ['🛡 경조사', '본인·가족 사유로 영업 부재'])
    add_para(doc, '💡 admin 승인 후 해당 월에만 적용 — 면제 시 페널티 무시 + 오버라이드 100% 지급', size=10)

    add_heading(doc, '⑦ KPI 보너스', level=1)
    add_para(doc, '없음 (V5.1) — 향후 정책 변경 시 별도 공지', bold=True, color=RGBColor(0x7f,0x1d,0x1d))

    add_heading(doc, '⑧ 계산 예시', level=1)
    # 예시 1
    add_para(doc, '예시 1) 정상 지급 (의무 충족 + 팀 이익률 충족)', bold=True, size=11, color=RGBColor(0x15,0x80,0x3d))
    add_para(doc, '조건: 본인 영업 22건 / G2 (20P, 우수 5건) / 팀원 인센 합계 800,000원 / 팀 이익률 28%', size=10, color=RGBColor(0x64,0x74,0x8b))
    t4 = doc.add_table(rows=0, cols=2); t4.style = 'Light List Accent 4'
    add_table_row(t4, ['항목','금액'], header=True)
    add_table_row(t4, ['본인 기본급', f'{fmt(BASE_SAL)}원'])
    add_table_row(t4, [f'본인 인센티브 (20P × {fmt(RATES["2"])})', f'+ {fmt(20*RATES["2"])}원'])
    add_table_row(t4, [f'본인 보너스 (5건 × {fmt(BONUS_PER)})', f'+ {fmt(5*BONUS_PER)}원'])
    ov1 = int(800000 * MGR_RATE * 1.0)
    add_table_row(t4, [f'🏢 팀 오버라이드 (800,000 × {int(MGR_RATE*100)}% × 100%)', f'+ {fmt(ov1)}원'])
    add_table_row(t4, ['페이백 차감', '− 0'])
    total1 = BASE_SAL + 20*RATES['2'] + 5*BONUS_PER + ov1
    add_table_row(t4, ['총 수령액', f'{fmt(total1)}원'], bold_first=True)

    # 예시 2
    add_para(doc, '\n예시 2) 부분 페널티 (본인 영업 12건, 의무 미달)', bold=True, size=11, color=RGBColor(0xca,0x8a,0x04))
    add_para(doc, '조건: 본인 영업 12건 (10건 이상 충족, 20건 미달) / 팀원 인센 800,000원 / 팀 이익률 28%', size=10, color=RGBColor(0x64,0x74,0x8b))
    t5 = doc.add_table(rows=0, cols=2); t5.style = 'Light List Accent 4'
    add_table_row(t5, ['항목','금액'], header=True)
    add_table_row(t5, ['본인 기본급+인센티브', '3,450,000원'])
    ov2 = int(800000 * MGR_RATE * 0.5)
    add_table_row(t5, [f'🏢 팀 오버라이드 (800,000 × {int(MGR_RATE*100)}% × 50%)', f'+ {fmt(ov2)}원'])
    add_table_row(t5, ['총 수령액', f'{fmt(3450000+ov2)}원'], bold_first=True)
    add_para(doc, '⚠️ 부분 페널티로 오버라이드 절반 차감', size=10, color=RGBColor(0xdc,0x26,0x26))

    # 예시 3
    add_para(doc, '\n예시 3) 미지급 (본인 영업 5건)', bold=True, size=11, color=RGBColor(0xdc,0x26,0x26))
    add_para(doc, '조건: 본인 영업 5건 (10건 미만) / 팀원 인센 1,000,000원', size=10, color=RGBColor(0x64,0x74,0x8b))
    t6 = doc.add_table(rows=0, cols=2); t6.style = 'Light List Accent 4'
    add_table_row(t6, ['항목','금액'], header=True)
    add_table_row(t6, ['본인 기본급+인센티브', '3,000,000원'])
    add_table_row(t6, [f'🏢 팀 오버라이드 (1,000,000 × {int(MGR_RATE*100)}% × 0%)', '+ 0원'])
    add_table_row(t6, ['총 수령액', '3,000,000원'], bold_first=True)
    add_para(doc, '🚨 본인 영업 10건 미만 — 오버라이드 전액 미지급', size=10, color=RGBColor(0xdc,0x26,0x26))

    add_heading(doc, '⑨ 팀장 행동 가이드', level=1)
    for line in [
        f'• 본인 영업 의무({MGR_OBLIG}건) 우선 달성 → 페널티 회피',
        f'• 팀원 매출 ↑ + 페이백 ↓ 코칭으로 팀 이익률 {int(MGR_PRMIN*100)}% 이상 유지',
        '• 본인 영업 부족 시 부분 페널티 (50% 차감) → 오버라이드 손실 큼',
        '• 팀원 우수 등록 독려 → 팀원 인센 합계 ↑ → 본인 오버라이드 ↑',
        '• 면제 사유 발생 시 admin에 즉시 보고 → 해당 월 페널티 무시',
    ]: add_para(doc, line, size=10)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(f'\n봉이모바일 팀장 인센티브 정책 {VER} · 마지막 업데이트 {EFFECT}')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64,0x74,0x8b); r.italic = True

    out = '/Users/vinsenzo/bongi-mobile/docs/incentive-doc-manager.docx'
    doc.save(out)
    print(f'✅ {out}')

if __name__ == '__main__':
    print(f'정책: {VER} (시행일 {EFFECT})')
    print(f'기본급 {fmt(BASE_SAL)}원 / G1 {fmt(RATES["1"])} / G2 {fmt(RATES["2"])} / G3 {fmt(RATES["3"])}')
    print(f'V5.1 manager: 오버라이드 {int(MGR_RATE*100)}% / 의무 {MGR_OBLIG}건 / 부분 {MGR_PARTIAL}건 / 팀이익률 {int(MGR_PRMIN*100)}%')
    print('---')
    gen_agent()
    gen_manager()
